"""DOCX parsing utilities"""
import logging
from typing import List, Dict, Any
from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph
from domain.models import DocumentOutline, DocumentSection

logger = logging.getLogger(__name__)


def extract_document_structure(file_path: str) -> DocumentOutline:
    """
    Extract complete structure from DOCX file
    Returns: DocumentOutline with sections, tables, metadata
    """
    try:
        doc = Document(file_path)
        
        sections = []
        word_count = 0
        tables_count = len(doc.tables)
        has_toc = False
        
        # Extract sections from paragraphs
        section_counter = {"1": 0, "2": 0, "3": 0, "4": 0}
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
                
            word_count += len(text.split())
            
            # Detect heading level
            level = 0
            if para.style.name.startswith('Heading'):
                try:
                    level = int(para.style.name.replace('Heading ', ''))
                except:
                    level = 0
            
            # Check for TOC
            if 'TOC' in para.style.name or 'tabla de contenido' in text.lower():
                has_toc = True
            
            if level > 0:
                # Update section numbering
                section_counter[str(level)] += 1
                for deeper_level in range(level + 1, 5):
                    section_counter[str(deeper_level)] = 0
                
                location = ".".join([str(section_counter[str(l)]) 
                                   for l in range(1, level + 1) 
                                   if section_counter[str(l)] > 0])
                
                sections.append(DocumentSection(
                    title=text,
                    level=level,
                    content=text,
                    location=f"Section {location}"
                ))
            elif len(sections) > 0:
                # Append content to last section
                sections[-1].content += f"\n{text}"
        
        # Extract metadata
        metadata = {
            "author": doc.core_properties.author or "",
            "created": str(doc.core_properties.created) if doc.core_properties.created else "",
            "modified": str(doc.core_properties.modified) if doc.core_properties.modified else "",
            "title": doc.core_properties.title or "",
            "subject": doc.core_properties.subject or "",
        }
        
        return DocumentOutline(
            filename=file_path.split("/")[-1],
            word_count=word_count,
            sections=sections,
            tables_count=tables_count,
            has_toc=has_toc,
            metadata=metadata
        )
        
    except Exception as e:
        logger.error(f"Error parsing DOCX: {e}", extra={"error": str(e)})
        raise


def search_in_document(outline: DocumentOutline, keywords: List[str]) -> List[Dict[str, str]]:
    """
    Search for keywords in document and return evidence
    Returns: [{"location": "Section X", "snippet": "..."}]
    """
    evidence = []
    
    for section in outline.sections:
        content_lower = section.content.lower()
        
        for keyword in keywords:
            if keyword.lower() in content_lower:
                # Extract snippet (50 chars before and after)
                idx = content_lower.find(keyword.lower())
                start = max(0, idx - 50)
                end = min(len(section.content), idx + len(keyword) + 50)
                snippet = section.content[start:end].strip()
                
                evidence.append({
                    "location": section.location,
                    "snippet": f"...{snippet}...",
                    "keyword": keyword
                })
    
    return evidence
